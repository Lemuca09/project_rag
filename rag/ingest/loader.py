"""Carrega arquivos suportados em documentos de texto."""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

SUPPORTED = {".txt", ".md", ".pdf", ".docx"}


@dataclass
class Document:
    source: str
    text: str
    metadata: dict = field(default_factory=dict)


def load_path(path: Path) -> list[Document]:
    path = path.resolve()
    if path.is_file():
        doc = _load_file(path)
        return [doc] if doc and doc.text.strip() else []
    if not path.is_dir():
        raise FileNotFoundError(f"Caminho inexistente: {path}")
    documents: list[Document] = []
    for file in sorted(path.rglob("*")):
        if file.is_file() and file.suffix.lower() in SUPPORTED:
            doc = _load_file(file)
            if doc and doc.text.strip():
                documents.append(doc)
    return documents


def _load_file(path: Path) -> Document | None:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return Document(source=str(path), text=text, metadata={"ext": suffix})
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".docx":
        return _load_docx(path)
    return None


def _load_pdf(path: Path) -> Document:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        content = page.extract_text() or ""
        if content.strip():
            pages.append(f"[pág. {i + 1}]\n{content}")
    return Document(
        source=str(path),
        text="\n\n".join(pages),
        metadata={"ext": ".pdf", "pages": len(reader.pages)},
    )


def _load_docx(path: Path) -> Document:
    import docx

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    return Document(
        source=str(path),
        text="\n".join(parts),
        metadata={"ext": ".docx"},
    )
